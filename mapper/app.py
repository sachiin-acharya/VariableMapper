from flask import Flask, request, send_file
from docx import Document
import os
from docx.shared import Pt
from docx.oxml import OxmlElement
import re
import json

app = Flask(__name__)


@app.route('/')
def index():
    return "Upload your JSON data to generate a Promissory Note."


def insert_customer_clause(paragraph, details):
    clause = f"lnlvtd {details.get('GrandFatherName', '')} ÷ {details.get('GrandMotherName', '')}, " \
             f"{details.get('FatherName', '')} ÷ {details.get('MotherName', '')}, {details.get('FatherInLawName', '')}, " \
             f"{details.get('SpouseName', '')}, {details.get('PermProvince', '')} {details.get('PermDistrict', '')}, " \
             f"{details.get('PermLocalLevelName', '')} j8f {details.get('PermWardNum', '')}. " \
             f"Temporary address: {details.get('TempProvince', '')}, {details.get('TempDistrict', '')}, " \
             f"{details.get('TempLocalLevelName', '')}, j8f {details.get('TempWardNum', '')}."
    paragraph.text = clause
    for run in paragraph.runs:
        run.font.name = 'Preeti'
        run.font.size = Pt(16)


def replace_placeholders(doc, data):
    for paragraph in doc.paragraphs:
        for placeholder, value in data.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in data.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, str(value))

# def set_cell_borders(table):
#     for row in table.rows:
#         for cell in row.cells:
#             # Create a new 'w:tcPr' element if it doesn't exist
#             # tc_pr = cell._element.get_or_add(OxmlElement('w:tcPr'))
#             tc_pr = cell._element.find('.//w:tcPr')
#             if tc_pr is None:
#                 tc_pr = OxmlElement('w:tcPr')
#                 cell._element.append(tc_pr)
#             # Create 'w:tcBorders' element
#             borders = OxmlElement('w:tcBorders')
#             # Define border properties
#             for border_type in ['top', 'left', 'bottom', 'right']:
#                 border = OxmlElement(f'w:{border_type}')
#                 border.set('w:val', 'single')  # Border style
#                 border.set('w:sz', '4')  # Size of the border
#                 border.set('w:space', '0')  # Space around the border
#                 borders.append(border)
#             # Append borders to the cell properties
#             tc_pr.append(borders)


@app.route('/generate', methods=['POST'])
def generate():
    json_data = request.get_json()
    customer_details = json_data['Data']['CustomerDetails']
    loan_details = customer_details.get('LoanDetails', [])
    authorized_persons = customer_details.get('AuthorizedPersonDetails', [])
    guarantor_details = customer_details.get('GuarantorDetails', [])
    branch_details = json_data['Data'].get('BranchDetails', {})

    doc = Document('PN Template/PromissoryNote.docx')

    placeholders_data = {
        '{{TotalLoanAmount}}': loan_details[0].get('LoanAmount', '0') if loan_details else '0',
        '{{LoanAmountInWords}}': loan_details[0].get('LoanAmountInWords', 'Zero') if loan_details else 'Zero',
        '{{GrandFatherName}}': customer_details.get('GrandFatherName', ''),
        '{{GrandMotherName}}': customer_details.get('GrandMotherName', ''),
        '{{FatherName}}': customer_details.get('FatherName', ''),
        '{{MotherName}}': customer_details.get('MotherName', ''),
        '{{FatherInLawName}}': customer_details.get('FatherInLawName', ''),
        '{{SpouseName}}': customer_details.get('SpouseName', ''),
        '{{PermProvince}}': customer_details.get('PermProvince', ''),
        '{{PermDistrict}}': customer_details.get('PermDistrict', ''),
        '{{PermLocalLevelName}}': customer_details.get('PermLocalLevelName', ''),
        '{{PermWardNo}}': customer_details.get('PermWardNum', ''),
        '{{TempProvince}}': customer_details.get('TempProvince', ''),
        '{{TempDistrict}}': customer_details.get('TempDistrict', ''),
        '{{TempLocalLevelName}}': customer_details.get('TempLocalLevelName', ''),
        '{{TempWardNo}}': customer_details.get('TempWardNum', ''),
        '{{Age}}': customer_details.get('Age', ''),
        '{{NameOfBorrower}}': customer_details.get('Name', ''),
        '{{CitizenshipNo}}': customer_details.get('CitizenshipNo', ''),
        '{{CitizenshipIssueDate}}': customer_details.get('CitizenshipIssueDate', ''),
        '{{CitizenshipIssueDistrict}}': customer_details.get('CitizenshipIssueDistrict', ''),
        '{{IssuingAuthority}}': customer_details.get('IssuingAuthority', ''),
        '{{RegProvince}}': branch_details.get('Province', ''),
        '{{RegDistrict}}': branch_details.get('District', ''),
        '{{RegLocalLevelName}}': branch_details.get('LocalBody', ''),
        '{{RegWardNo}}': branch_details.get('WardNum', ''),
        '{{DesignationOfAuthorizedPerson}}': authorized_persons[0].get('Designation', '') if authorized_persons else '',
        '{{LoanID}}': loan_details[0].get('LoanId', 'NoLoanID') if loan_details else 'NoLoanID',
        '{{FacilityName}}': loan_details[0].get('FacilityName', '') if loan_details else '',
        '{{InterestRate}}': loan_details[0].get('InterestRate', '') if loan_details else '',
        '{{NameOfPersonalGuarantors}}': ', '.join(
            guarantor['Name'] for guarantor in guarantor_details) if guarantor_details else ''
    }

    table = doc.tables[1]

    if loan_details:
        first_loan = loan_details[0]
        row_cells = table.row_cells(1)
        row_cells[0].text = first_loan["LoanId"]
        row_cells[1].text = first_loan["LoanName"]
        row_cells[2].text = f"{first_loan['LoanAmount']}--cIf/]kL === {first_loan['LoanAmountInWords']} dfq_"
        row_cells[3].text = f"{first_loan['InterestRate']}k|ltzt"

    for loan in loan_details[1:]:
        row_cells = table.add_row().cells
        row_cells[0].text = loan["LoanId"]
        row_cells[1].text = loan["LoanName"]
        row_cells[2].text = f"{loan['LoanAmount']}--cIf/]kL === {loan['LoanAmountInWords']} dfq_"
        row_cells[3].text = f"{loan['InterestRate']}k|ltzt"

    # set_cell_borders(table)


    if isinstance(customer_details, list):
        for customer in customer_details:
            paragraph = doc.add_paragraph()
            insert_customer_clause(paragraph, customer)

    if customer_details.get('CustomerType') != 'Institutional':
        block_to_remove = "{{IssuingAuthority}} df btf{ /x]sf] {{RegProvince}} k|b]z {{RegDistrict}} " \
                          "lhNnf, =========={{RegLocalLevelName}} d=g=kf=÷p=d=g=kf=÷g=kf=÷÷uf=kf= j8f g+= {{RegWardNo}} " \
                          "l:yt /lhi68{ sfof{no /x]sf] {{NameOfBorrower}} -h;nfO{ o;kl5 C0fL sDkgL÷kmd{ elgPsf] 5_ sf " \
                          "tkm{af6 clVtof/ k|fKt {{DesignationOfAuthorizedPerson}} kbsf {{GrandFatherName}} ÷{{GrandMotherName}}  " \
                          "sf]÷sL] gflt÷gfltgL, {{FatherName}} ÷ {{MotherName}}  sf]÷sL 5f]/f÷5f]/L, {{FatherInLawName}} sf] a'xf/L, " \
                          "{{SpouseName}} sf] klt÷klTg, {{PermProvince}} k|b]z {{PermDistrict}} lhNnf, {{PermLocalLevelName}} " \
                          "d=g=kf=÷p=d=g=kf=÷g=kf÷uf=kf=, j8f g+= {{PermWardNo}}:yfoL 7]ufgf eO{ xfn {{TempProvince}} k|b]z {{TempDistrict}} lhNnf, " \
                          "{{TempLocalLevelName}} d=g=kf=÷p=d=g=kf=÷g=kf÷uf=kf= j8f g+ {{TempWardNo}} a:g] jif{ {{Age}} sf]÷sL] {{NameOfBorrower}} -gf=k|=g+= " \
                          "{{CitizenshipNo}}, hf/L ldlt {{CitizenshipIssueDate}}, lhNnf k|zf;g sfof{no {{CitizenshipIssueDistrict}}_  "
        company_stamp = "sDkgLsf] 5fkM"
        for paragraph in doc.paragraphs:
            if block_to_remove in paragraph.text:
                paragraph.text = paragraph.text.replace(block_to_remove, '')
            if company_stamp in paragraph.text:
                paragraph.text = paragraph.text.replace(company_stamp, '')

                for authorized_person in authorized_persons:
                    clause = f"{{{{DesignationOfAuthorizedPerson}}}} kbsf {authorized_person.get('GrandFatherName', '')} ÷ {authorized_person.get('GrandMotherName', '')} sf]÷sL] gflt÷gfltgL, " \
                             f"{authorized_person.get('FatherName', '')} ÷ {authorized_person.get('MotherName', '')} sf]÷sL 5f]/f÷5f]/L, " \
                             f"{authorized_person.get('FatherInLawName', '')} sf] a'xf/L, {authorized_person.get('SpouseName', '')} sf] klt÷klTg, " \
                             f"{authorized_person.get('PermProvince', '')} k|b]z {authorized_person.get('PermDistrict', '')} lhNnf, " \
                             f"{authorized_person.get('PermLocalLevelName', '')} d=g=kf=÷p=d=g=kf=÷g=kf÷uf=kf=, j8f g+= " \
                             f"{{{{PermWardNo}}}}:yfoL 7]ufgf eO{{ xfn {authorized_person.get('TempProvince', '')} k|b]z {authorized_person.get('TempDistrict', '')} lhNnf, " \
                             f"{authorized_person.get('TempLocalLevelName', '')} d=g=kf=÷p=d=g=kf=÷g=kf÷uf=kf= j8f g+ {{TempWardNo}} a:g] " \
                             f"jif{{ {authorized_person.get('Age', '')} sf]÷sL] {authorized_person.get('NameOfBorrower', '')} -gf=k|=g+= " \
                             f"{{CitizenshipNo}}, hf/L ldlt {{CitizenshipIssueDate}}, lhNnf k|zf;g sfof{{no {{CitizenshipIssueDistrict}}_"
                    new_paragraph = doc.add_paragraph(clause)
                    for run in new_paragraph.runs:
                        run.font.name = 'Preeti'
                        run.font.size = Pt(16)

    replace_placeholders(doc, placeholders_data)

    if customer_details.get('MaritalStatus') == 'Married':
        if customer_details.get('Gender') == 'Female':
            doc.add_paragraph(f"{customer_details.get('FatherInLawName', '')} sf] a'xf/L,")
        doc.add_paragraph(
            f"{customer_details.get('FatherInLawName', '')} sf] a'xf/L, {customer_details.get('SpouseName', '')} sf] klt÷klTg.")


    if customer_details.get('CustomerType') == 'Institutional':
        paragraph = doc.add_paragraph()
        paragraph.text = f"{customer_details.get('IssuingAuthority', '')} issued in {branch_details.get('Province', '')}, " \
                         f"{branch_details.get('District', '')}, {branch_details.get('LocalBody', '')} Ward {branch_details.get('WardNum', '')}."
        for run in paragraph.runs:
            run.font.name = 'Preeti'
            run.font.size = Pt(16)

    if loan_details:
        for loan in loan_details:
            loan_clause = f"s'n shf{{ l;df ¿= {loan.get('LoanAmount', '0')} - {loan.get('LoanAmountInWords', 'Zero')}, " \
                          f"Guaranteed by {guarantor_details[0].get('Name', '')} (if applicable)."
            doc.add_paragraph(loan_clause)

    output_folder = './Output'
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    borrower_name = customer_details[0].get('Name', 'Unknown').replace(' ', '_') if isinstance(customer_details,list) else customer_details.get('Name', 'Unknown').replace(' ', '_')
    loan_id = loan_details[0].get('LoanId', 'NoLoanID') if loan_details else 'NoLoanID'
    output_file_path = os.path.join(output_folder, f"Promissory_Note_{borrower_name}_{loan_id}.docx")

    doc.save(output_file_path)

    return send_file(output_file_path, as_attachment=True)


if __name__ == '__main__':
    app.run(debug=True)
